import io
from pathlib import Path
from docx import Document


TEMPLATE_PATH = Path(__file__).resolve().parent.parent / 'templates' / 'plantilla_factura.docx'

CONCEPTO_DISPLAY = {
    'venta': 'Venta',
    'separacion': 'Separación',
}
ITEM_DISPLAY = {
    'laptop': 'Laptop',
    'tarjeta_grafica': 'Tarjeta Gráfica',
    'hardware': 'Hardware',
    'pc_mesa': 'PC de Mesa',
}
PAYMENT_DISPLAY = {
    'efectivo': 'Efectivo',
    'tarjeta': 'Tarjeta',
    'transferencia': 'Transferencia',
    'otro': 'Otro',
}


def _format_currency(amount):
    try:
        int_amount = int(amount)
        formatted = f"{int_amount:,}".replace(',', '.')
        return f"COP ${formatted}"
    except (ValueError, TypeError):
        return str(amount)


def _replace_in_paragraph(paragraph, placeholders):
    for ph, value in placeholders.items():
        if ph in paragraph.text:
            for run in paragraph.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, value)
                    break
            else:
                full_text = paragraph.text.replace(ph, value)
                for i, run in enumerate(paragraph.runs):
                    run.text = full_text if i == 0 else ''


def generate_invoice_document(invoice) -> io.BytesIO:
    """
    Load the .docx template, fill all placeholders, and return a BytesIO
    with the generated Word document.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found at: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    placeholders = {
        '{{concepto}}': CONCEPTO_DISPLAY.get(invoice.concepto, invoice.concepto),
        '{{item}}': ITEM_DISPLAY.get(invoice.item, invoice.item),
        '{{serial_item}}': str(invoice.serial_item),
        '{{total_amount}}': _format_currency(invoice.total_amount),
        '{{due_date}}': invoice.due_date.strftime('%d/%m/%Y'),
        '{{payment_method}}': PAYMENT_DISPLAY.get(invoice.payment_method, invoice.payment_method),
        '{{bill_id}}': str(invoice.bill_id),
        '{{client_name}}': str(invoice.client_name),
        '{{client_id}}': str(invoice.client_document),
        '{{client_phone}}': str(invoice.client_phone),
        '{{client_address}}': str(invoice.client_address),
        '{{client_email}}': str(invoice.client_email),
    }

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, placeholders)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, placeholders)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
